import io

import pytest

from secretary_bot import documents


def test_extract_plain_text():
    assert documents.extract_text("notes.md", "решили X".encode("utf-8")) == "решили X"
    assert documents.extract_text("notes.txt", "привет".encode("cp1251")) == "привет"


def test_extract_binary_text_rejected():
    with pytest.raises(ValueError):
        documents.extract_text("data.bin", b"\x00\x01\x02binary")


def test_old_doc_rejected_with_hint():
    with pytest.raises(ValueError):
        documents.extract_text("report.doc", b"\xd0\xcf\x11\xe0 ole header")


def test_extract_docx_round_trip():
    from docx import Document

    d = Document()
    d.add_paragraph("Решили использовать SQLite")
    d.add_paragraph("Аргумент: проще эксплуатация")
    buf = io.BytesIO()
    d.save(buf)
    text = documents.extract_text("report.docx", buf.getvalue())
    assert "SQLite" in text and "Аргумент" in text
